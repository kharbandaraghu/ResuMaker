from docx import Document, oxml, opc
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.text.run import Run
from docx.enum.section import WD_ORIENT
from typing import List

class TextStyle:
    def __init__(self, isHyperlink=False, url='', font_size=12, font_name='GoogleSans', Bold=False, Italic=False, underline=False, font_color=(32,32,32)):
        self.isHyperlink = isHyperlink
        self.url = url
        self.font_size = font_size
        self.font_name = font_name
        self.Bold = Bold
        self.Italic = Italic
        self.underline = underline
        self.font_color = font_color

class ParagraphStyle:
    def __init__(self, style='Normal', align='left', space_before=0, space_after=0, indent_level=0, line_spacing=1.0):
        self.style = style
        self.align = align
        self.space_before = space_before
        self.space_after = space_after
        self.indent_level = indent_level
        self.line_spacing = line_spacing

class Run:
    def __init__(self, value: str, style: TextStyle):
        self.value = value
        self.style = style

class ResumeMaker:
    def __init__(self, margins=[0.5, 0.5, 0.5, 0.5], height=11, width=8.5):
        self.doc = Document()
        self.doc.sections[0].page_height = Inches(height)
        self.doc.sections[0].page_width = Inches(width)
        self.doc.sections[0].left_margin = Inches(margins[0])
        self.doc.sections[0].right_margin = Inches(margins[1])
        self.doc.sections[0].top_margin = Inches(margins[2])
        self.doc.sections[0].bottom_margin = Inches(margins[3])
        self.doc.sections[0].orientation = WD_ORIENT.PORTRAIT

    def save(self, path):
        self.doc.save(path)

    def __addRun(self, paragraph, run):
        text_run = paragraph.add_run(run.value)
        text_run.font.size = Pt(run.style.font_size)
        text_run.font.name = run.style.font_name
        text_run.font.bold = run.style.Bold
        text_run.font.italic = run.style.Italic
        text_run.font.color.rgb = RGBColor(run.style.font_color[0], run.style.font_color[1], run.style.font_color[2])
    
    def addImage(self, path, width=1.5, height=1.5):
        self.doc.add_picture(path, width=Inches(width), height=Inches(height))
    
    def addMultipleText(self, runs: List[Run], paragraphStyle: ParagraphStyle):
        paragraph = self.doc.add_paragraph(style=paragraphStyle.style)
        align_mapping = {'left': WD_ALIGN_PARAGRAPH.JUSTIFY, 'center': WD_ALIGN_PARAGRAPH.CENTER, 'right': WD_ALIGN_PARAGRAPH.RIGHT}
        align = align_mapping.get(paragraphStyle.align, WD_ALIGN_PARAGRAPH.LEFT)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing = paragraphStyle.line_spacing
        paragraph_format.space_before = Pt(paragraphStyle.space_before)
        paragraph_format.space_after = Pt(paragraphStyle.space_after)
        paragraph_format.alignment = align
        paragraph_format.left_indent = Inches(0.5 * paragraphStyle.indent_level)
        for run in runs:
            if run.style.isHyperlink:
                self.__addHyperlink(paragraph, run)
                # add space
                self.__addRun(paragraph, Run(' ', run.style))
            else:
                self.__addRun(paragraph, run)
                self.__addRun(paragraph, Run(' ', run.style))
    
    def addLeftRightText(self, left: Run, right: Run, paragraphStyle: ParagraphStyle):
        paragraph = self.doc.add_paragraph(style=paragraphStyle.style)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing = paragraphStyle.line_spacing
        paragraph_format.space_before = Pt(paragraphStyle.space_before)
        paragraph_format.space_after = Pt(paragraphStyle.space_after)
        paragraph_format.left_indent = Inches(0.5 * paragraphStyle.indent_level)

        # Add left run
        self.__addRun(paragraph, left)
        # Add tab stop to align right run
        paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(self.doc.sections[0].page_width.inches - self.doc.sections[0].left_margin.inches - self.doc.sections[0].right_margin.inches), WD_TAB_ALIGNMENT.RIGHT)

        # Add tab character
        self.__addRun(paragraph, Run('\t', TextStyle()))

        # Add right run
        self.__addRun(paragraph, right)

    def addHyperlink(self, text, url, style: TextStyle, paragraphStyle: ParagraphStyle):
        paragraph = self.doc.add_paragraph(style=paragraphStyle.style)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing = paragraphStyle.line_spacing
        paragraph_format.space_before = Pt(paragraphStyle.space_before)
        paragraph_format.space_after = Pt(paragraphStyle.space_after)
        paragraph_format.alignment = paragraphStyle.align
        paragraph_format.left_indent = Inches(0.5 * paragraphStyle.indent_level)

        run = Run(text, style)
        self.__addHyperlink(paragraph, run)
        
    def __addHyperlink(self, paragraph, run):
        part = paragraph.part
        r_id = part.relate_to(run.style.url, opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        hyperlink = oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(oxml.shared.qn('r:id'), r_id)

        new_run = oxml.shared.OxmlElement('w:r')
        rPr = oxml.shared.OxmlElement('w:rPr')

        if run.style.font_color is not None:
            c = oxml.shared.OxmlElement('w:color')
            c.set(oxml.shared.qn('w:val'), '%02x%02x%02x' % run.style.font_color)
            rPr.append(c)
        
        if run.style.font_name is not None:
            rFonts = oxml.shared.OxmlElement('w:rFonts')
            rFonts.set(oxml.shared.qn('w:ascii'), run.style.font_name)
            rFonts.set(oxml.shared.qn('w:cs'), run.style.font_name)
            rFonts.set(oxml.shared.qn('w:eastAsia'), run.style.font_name)
            rFonts.set(oxml.shared.qn('w:hAnsi'), run.style.font_name)
            rPr.append(rFonts)
        
        if run.style.Bold:
            b = oxml.shared.OxmlElement('w:b')
            rPr.append(b)

        if run.style.Italic:
            i = oxml.shared.OxmlElement('w:i')
            rPr.append(i)
        
        sz = oxml.shared.OxmlElement('w:sz')
        sz.set(oxml.shared.qn('w:val'), str(run.style.font_size*2))
        rPr.append(sz)

        if run.style.underline:
            u = oxml.shared.OxmlElement('w:u')
            u.set(oxml.shared.qn('w:val'), 'single')
            rPr.append(u)

        new_run.append(rPr)
        new_run.text = run.value
        hyperlink.append(new_run)

        paragraph._p.append(hyperlink)

        return hyperlink